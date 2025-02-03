import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import streamlit as st


def generate_word_document(data, output_folder='output'):
    """Generates personalized Word documents for each row in the data."""
    for index, row in data.iterrows():
        doc = Document()
        doc.add_heading(f'Employee Information: {row["Name"]}', level=1)
        doc.add_paragraph(f'Email: {row["Email"]}')
        doc.add_paragraph(f'Company Name: {row["Company Name"]}')
        doc.add_paragraph(f'Position: {row["Position"]}')
        doc.add_paragraph(f'Joining Date: {row["Joining Date"]}')
        document_name = f'{row["Name"]}_{row["Company Name"]}.docx'
        document_path = f'{output_folder}/{document_name}'
        doc.save(document_path)


def generate_pdf_document(data, output_folder='output'):
    """Generates personalized PDF documents for each row in the data."""
    for index, row in data.iterrows():
        c = canvas.Canvas(
            f'{output_folder}/{row["Name"]}_{row["Company Name"]}.pdf',
            pagesize=letter)
        c.drawString(inch, 10 * inch, f'Employee Information: {row["Name"]}')
        c.drawString(inch, 9 * inch, f'Email: {row["Email"]}')
        c.drawString(inch, 8 * inch, f'Company Name: {row["Company Name"]}')
        c.drawString(inch, 7 * inch, f'Position: {row["Position"]}')
        c.drawString(inch, 6 * inch, f'Joining Date: {row["Joining Date"]}')
        c.save()


def main():
    st.title("Personalized Document Generator")

    input_method = st.radio("Choose input method:",
                            ("Excel Upload", "Manual Entry"))

    if input_method == "Excel Upload":
        st.write("## Excel Upload")
        st.write(
            "Choose the 'Upload Excel File' option. Upload one or more Excel files containing the required columns:"
        )
        st.write("- Name")
        st.write("- Email")
        st.write("- Company Name")
        st.write("- Position")
        st.write("- Joining Date")
        uploaded_files = st.file_uploader("Upload one or more Excel files",
                                          type=["xlsx", "xls"],
                                          accept_multiple_files=True)
        if uploaded_files:
            dataframes = []
            for uploaded_file in uploaded_files:
                df = pd.read_excel(uploaded_file)
                dataframes.append(df)
            data = pd.concat(dataframes, ignore_index=True)
            st.dataframe(data)

    elif input_method == "Manual Entry":
        st.write("## Manual Entry")
        st.write(
            "Select the 'Manual Entry' option to add employee details one by one. Use the form to fill in the data and add entries to the list. When ready, click Generate Documents to create the documents."
        )
        name = st.text_input("Name")
        email = st.text_input("Email")
        company_name = st.text_input("Company Name")
        position = st.text_input("Position")
        joining_date = st.date_input("Joining Date")
        data = pd.DataFrame(
            [[name, email, company_name, position, joining_date]],
            columns=[
                "Name", "Email", "Company Name", "Position", "Joining Date"
            ])
        st.dataframe(data)

    if st.button("Generate Documents"):
        output_folder = 'output'
        import os
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        generate_word_document(data, output_folder)
        generate_pdf_document(data, output_folder)
        st.success("Documents generated successfully!")


if __name__ == "__main__":
    main()
